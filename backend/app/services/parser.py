import io
from typing import List, Tuple


class DocumentParser:
    SUPPORTED_TYPES = {"pdf", "docx", "doc", "html", "htm", "txt"}

    @staticmethod
    def parse(file_bytes: bytes, filename: str) -> Tuple[str, str]:
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
        if ext not in DocumentParser.SUPPORTED_TYPES:
            raise ValueError(f"不支持的文件类型: .{ext}")

        if ext == "pdf":
            return DocumentParser._parse_pdf(file_bytes, filename)
        elif ext in ("docx", "doc"):
            return DocumentParser._parse_docx(file_bytes, filename)
        elif ext in ("html", "htm"):
            return DocumentParser._parse_html(file_bytes, filename)
        elif ext == "txt":
            return DocumentParser._parse_txt(file_bytes, filename)
        else:
            raise ValueError(f"不支持的文件类型: .{ext}")

    @staticmethod
    def _parse_pdf(file_bytes: bytes, filename: str) -> Tuple[str, str]:
        import pdfplumber

        texts: List[str] = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text:
                    header = f"[第{i + 1}页] "
                    texts.append(header + page_text)

        full_text = "\n\n".join(texts)
        title = filename.rsplit(".", 1)[0]
        return full_text, title

    @staticmethod
    def _parse_docx(file_bytes: bytes, filename: str) -> Tuple[str, str]:
        from docx import Document as DocxDocument

        doc = DocxDocument(io.BytesIO(file_bytes))
        paragraphs: List[str] = []

        for para in doc.paragraphs:
            style = para.style.name if para.style else ""
            text = para.text.strip()
            if not text:
                continue
            if "Heading" in style or "heading" in style or "标题" in style:
                paragraphs.append(f"【{text}】")
            else:
                paragraphs.append(text)

        for table in doc.tables:
            table_lines: List[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                table_lines.append(" | ".join(cells))
            if table_lines:
                paragraphs.append("\n" + "\n".join(table_lines) + "\n")

        full_text = "\n\n".join(paragraphs)
        title = filename.rsplit(".", 1)[0]
        return full_text, title

    @staticmethod
    def _parse_html(file_bytes: bytes, filename: str) -> Tuple[str, str]:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(file_bytes, "lxml")

        for tag in soup(["script", "style", "nav", "footer", "header", "aside", "noscript"]):
            tag.decompose()

        title_tag = soup.find("title")
        title = title_tag.get_text(strip=True) if title_tag else filename.rsplit(".", 1)[0]

        body = soup.find("body")
        if body:
            text = body.get_text(separator="\n", strip=True)
        else:
            text = soup.get_text(separator="\n", strip=True)

        return text, title

    @staticmethod
    def _parse_txt(file_bytes: bytes, filename: str) -> Tuple[str, str]:
        text = file_bytes.decode("utf-8", errors="replace")
        title = filename.rsplit(".", 1)[0]
        return text, title


document_parser = DocumentParser()
